from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT=Path(r"C:\Users\94012\Desktop\gaussian 参数.docx")
def font(r,size=10.2,bold=False,color="222222"):
    r.font.name="Microsoft YaHei"
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Microsoft YaHei")
    r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
def shade(c,fill):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:fill"),fill); p.append(s)
def putcell(c,txt,bold=False):
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.03
    font(p.add_run(str(txt)),9.1,bold,"17365D" if bold else "222222")
def table(doc,heads,rows,widths):
    t=doc.add_table(rows=1,cols=len(heads)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    for i,h in enumerate(heads):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]/1440); shade(c,"DCE6F1"); putcell(c,h,True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].width=Inches(widths[i]/1440); putcell(cs[i],v)
    return t
def para(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.08; font(p.add_run(txt)); return p
def bullets(doc,items):
    for x in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); font(p.add_run(x),10)

doc=Document(); s=doc.sections[0]
s.top_margin=Inches(.7); s.bottom_margin=Inches(.65); s.left_margin=Inches(.75); s.right_margin=Inches(.75)
for st,sz,col in [("Normal",10.5,"222222"),("Heading 1",16,"1F4E79"),("Heading 2",13,"2E75B6"),("Heading 3",11,"1F4E79")]:
    z=doc.styles[st]; z.font.name="Microsoft YaHei"; z._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); z.font.size=Pt(sz); z.font.color.rgb=RGBColor.from_string(col); z.font.bold=st!="Normal"
h=s.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(h.add_run("R²-Gaussian · SiO₂ 参数参考手册"),8.5,False,"6B7280")
f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(f.add_run("当前仓库参数说明 · 2026-09-02"),8,False,"6B7280")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("R²-Gaussian 项目参数参考手册"),22,True,"17365D")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("SiO₂：旋转中心 → 自定义 FBP → 数据划分 → 点云初始化 → 训练 → 测试"),11,False,"5B6573")
para(doc,"用途：集中记录当前项目所有主要命令行参数，解释参数是什么、有什么用、改变后会发生什么，并给出 SiO₂ 示例。你的数据为181张0°–180°、1°间隔的强度TIFF，尺寸约976×1028，按ParallelBeam处理。")

doc.add_heading("1. 总体流程与路径",1)
table(doc,["阶段","脚本","输出"],[("旋转中心","scripts/scan_fbp_center_tomopy.py / scan_fbp_center.py","center.json、center.csv或中心扫描CSV"),("FBP与划分","prepare_fbp_tiff.py","proj_all、proj_train、proj_test、vol_fbp、meta_data.json"),("点云初始化","init_from_fbp.py / initialize_pcd.py","init_*.npy"),("训练","train.py","模型、评估、checkpoint"),("测试","test.py","渲染、体重建、PSNR/SSIM")],[1800,4000,3760])
para(doc,"路径变量示例：RAW=/root/autodl-tmp/3D-GS/SiO2；DATA=/root/autodl-tmp/3D-GS/sio2_r2_fbp_128_v1/dataset；MODEL=/root/autodl-tmp/3D-GS/sio2_r2_fbp_128_v1/model_sio2_128。")

doc.add_heading("2. 旋转中心：scan_fbp_center_tomopy.py",1)
para(doc,"功能：读取原始TIFF，预处理后估计0°/180°旋转中心；输出的offDetector[0]会用于FBP。宽度1028时理论中心约514 px。")
center_rows=[
("input_dir","输入TIFF目录；会排除Flat平场图。"),("config","读取NumberImages=181、AngleFirst=0、AngleInterval=1；数量不符会报错。"),("angles_file","角度表第二列；角度不规则时使用。"),("angle_start / angle_interval","无config时的起始角度/间隔（度）。"),("keep_duplicate_endpoint","保留180°端点；旧版181 vs 180报错时可临时添加。"),("input_type","intensity执行-log(I/I0)；line_integral表示已是线积分。"),("i0","入射强度；0表示逐张自动估计。已知I0=14000时可写14000。"),("i0_percentile","自动I0的分位数；99.5可避免异常亮点决定I0。"),("zero_policy","nearest用邻近有效像素填零；clip截断；keep风险较高。"),("log_eps","取对数前最小正值，防止log(0)。"),("clip_percentile","NaN/Inf替换上限；99.9表示用99.9%分位数。"),("shift_v","上下方向平移像素；5是旧流程经验值，不确定时先0。"),("pixel_subsample","中心估计降采样倍数；1保留原始精度，4更快但精度下降。"),("resize_order","降采样插值阶数；1为线性，质量/速度平衡。"),("projection_scale","线积分整体乘数；1表示不缩放。"),("pixel_size","像素物理尺寸；偏移-3px且尺寸0.02时offDetector=-0.06。"),("method","vo用0°/180°对称性；scipy优化法可能发散。"),("init_px","scipy初始中心；1028宽图可设514。"),("tol","搜索步长；0.25粗搜，0.05精搜。"),("algorithm","传给TomoPy的算法名，通常保持scipy。"),("search_min_px / search_max_px","vo相对中心的搜索范围；-100到100表示大范围粗搜。"),("slice_step","检测器行采样间隔；16更快，8更密。"),("slice_margin","忽略上下边缘行数；256可避开边缘截断。"),("max_slices","最多使用检测器行数；多个行取中位数，0表示不限制。"),("threads","CPU线程数；服务器可设8–16。"),("output","结果JSON；父目录需先mkdir -p。")]
table(doc,["参数","解释、作用与例子"],center_rows,[2500,6860])
para(doc,"结果示例：center_px=516.5时，offset_px=514-516.5=-2.5，offDetector[0]=-2.5×0.02=-0.05。若不同检测器行的spread达几十像素或中心越界（如1523px），不要直接用于FBP。")

doc.add_heading("3. 备用中心扫描：scan_fbp_center.py",1)
para(doc,"功能：对处理后的0°/180°端点投影做有界相关性扫描。Tomopy结果异常时使用。先粗搜，再围绕最佳偏移精搜。")
table(doc,["参数","解释、作用与例子"],[("data","包含meta_data.json和proj_endpoint_pair.npy的临时数据目录。"),("u_min_px / u_max_px","水平偏移搜索范围；粗搜-100到100，已知约-3px后可精搜-5到-1。"),("u_step_px","步长；粗搜1px，精搜0.25或0.05px。"),("endpoint_pair","手动指定端点文件；默认从metadata读取。"),("margin_px","左右边缘保护像素，避免边界相关性干扰。"),("output","CSV输出；父目录必须存在。")],[2500,6860])

doc.add_heading("4. 自定义FBP与数据划分：prepare_fbp_tiff.py",1)
para(doc,"功能：TIFF→坏点处理→强度转线积分→中心偏移→TIGRE平行束FBP→体归一化→训练/测试划分。输出目录必须为空。")
fbp_rows=[("input_dir / output_dir","输入TIFF和全新输出目录。"),("config / angles_file / center_json","角度配置与旋转中心；center_json会覆盖offDetector。"),("input_type / i0 / i0_percentile","强度转线积分参数；原始强度用intensity、i0=0。"),("zero_policy / log_eps / clip_percentile","坏像素和数值安全处理。"),("shift_v","纵向补偿；中心估计和FBP必须保持一致。"),("pixel_subsample / resize_order","降采样倍数/插值；4约把976×1028变为244×244。"),("projection_scale","投影整体缩放；通常1。"),("volume_sign","auto在FBP整体为负时翻转。"),("volume_background_percentile / sigma","逐轴向扣背景及平滑；-1关闭，1/8适合缓慢基线。"),("volume_normalize_percentile","正值分位数归一化；99.5约映射到[0,1]，0关闭。"),("n_train / n_test","训练/测试视角数；唯一180视角可用120/60。"),("nVoxel / sVoxel","FBP体素数量/物理尺寸；128³适合首轮，256³更细但更慢。"),("offOrigin / offDetector","体积中心偏移/探测器[u,v]偏移；后者用中心结果。"),("DSD / DSO","源到探测器/旋转中心距离；保持项目几何约定。"),("pixel_size / accuracy","像素尺寸与TIGRE射线精度；accuracy越小通常越准但更慢。"),("filter","FBP滤波器；hann抑噪，ram_lak更锐但放大噪声。"),("seed","固定测试集随机抽样，保证复现。")]
table(doc,["参数","解释、作用与例子"],fbp_rows,[3000,6360])

doc.add_heading("5. FBP点云初始化：init_from_fbp.py",1)
para(doc,"功能：从vol_fbp.npy选择高密度体素，映射到[-1,1]³，生成每行[x,y,z,density]的初始Gaussian。")
table(doc,["参数","解释、作用与例子"],[("volume","输入vol_fbp.npy。"),("output","输出init_*.npy；自定义名称时还会保存canonical文件。"),("n_points","点数；50000是常用起点，增加会提高覆盖但占显存。"),("density_thresh","体素密度阈值；点不够时0.05→0.02。"),("density_rescale","初始密度缩放；0.15通常稳定，过大可能早期投影过强。"),("normalize_percentile","再次归一化；vol_fbp已归一化时设0避免重复。"),("seed","固定采样。"),("visualize / visualize_show","保存/打开三维点云预览；服务器只用visualize。"),("visualize_output / visualize_max_points","预览路径/最多绘制点数。")],[3000,6360])

doc.add_heading("6. 通用初始化：initialize_pcd.py",1)
para(doc,"适用于直接从投影用FBP/FDK重建初始化或随机初始化。已经有vol_fbp.npy的SiO₂优先用init_from_fbp.py。")
table(doc,["参数","解释、作用与例子"],[("data / output","数据集目录与输出npy；不写output则自动命名。"),("evaluate","用数据集体数据评估初始3D PSNR，仅调试使用。"),("recon_method","random/fbp/fdk；parallel应使用fbp。"),("recon_split","all使用proj_all全部视角；train只用训练视角。"),("n_points / density_thresh / density_rescale","点数、体素阈值、密度缩放。"),("random_density_max","random模式最大随机密度。")],[3000,6360])

doc.add_heading("7. 训练：train.py",1)
doc.add_heading("7.1 数据、模型与管线",2)
table(doc,["参数","解释、作用与例子"],[("CUDA_VISIBLE_DEVICES","GPU选择；0表示第0卡。"),("-s / source_path","数据集目录，读取meta_data.json和投影。"),("-m / model_path","模型、评估、日志、checkpoint输出目录。"),("--ply_path","初始化.npy/.ply；显式指定避免自动找错。"),("--scale_min / scale_max","Gaussian尺度上下限，占体积尺寸比例；过大限制细节。"),("--eval","读取测试集并评估；建议开启。"),("--data_device","数据设备，通常cuda。"),("--compute_cov3D_python","Python计算协方差；通常False更快。"),("--debug","底层调试输出。"),("--detect_anomaly","检测NaN/异常梯度；调试用，会变慢。"),("--quiet","减少终端输出。"),("--config","用配置文件覆盖命令行参数。")],[3000,6360])

doc.add_heading("7.2 学习率",2)
table(doc,["参数组","初始→最终","解释与调整"],[("position_lr_init/final/max_steps","0.0002→0.00002 / 30000","位置(x,y,z)更新；初始偏差大可略增，震荡则降低。"),("density_lr_init/final/max_steps","0.01→0.001 / 30000","密度/衰减贡献；拟合慢可增，波动大可减。"),("scaling_lr_init/final/max_steps","0.005→0.0005 / 30000","椭球大小更新；控制模糊与覆盖范围。"),("rotation_lr_init/final/max_steps","0.001→0.0001 / 30000","椭球方向更新；细长结构有用，过大可能抖动。")],[3000,2200,4160])
para(doc,"每组学习率按指数方式衰减。例：位置学习率从0.0002逐步降到0.00002，前期快速纠正FBP点位，后期稳定收敛。")
doc.add_heading("7.3 损失函数",2)
table(doc,["参数","解释、作用与例子"],[("lambda_dssim=0.25","结构相似损失权重；0关闭，增大更重视边缘和局部对比度。"),("lambda_tv=0.05","三维总变分正则；增大更平滑、噪声少，但可能抹细节。"),("tv_vol_size=32","每次TV随机小体积为32³；增大更全面但更耗显存。")],[3000,6360])
doc.add_heading("7.4 Gaussian自适应控制",2)
table(doc,["参数","解释、作用与例子"],[("density_min_threshold=1e-5","低密度Gaussian可删除；过高会删弱结构。"),("densification_interval=100","每100步检查增删Gaussian。"),("densify_from_iter=500","500步后开始增点，先稳定初始点云。"),("densify_until_iter=15000","15000步后停止增点，后半程只优化已有点；可改20000保细节。"),("densify_grad_threshold=5e-5","位置梯度阈值；越低越易复制/拆分，点数和显存增大。"),("densify_scale_threshold=0.1","尺度比例阈值；sVoxel=2时约对应0.2，用于clone/split。"),("max_screen_size / max_scale","2D/3D尺度裁剪上限；None表示不按该项裁剪。"),("max_num_gaussians=500000","Gaussian数量上限；显存不足可300000，细节不足且显存足可提高。")],[3000,6360])
doc.add_heading("7.5 训练、评估、保存和断点",2)
table(doc,["参数","解释、作用与例子"],[("iterations=30000","总训练步数；5000试跑，30000正式训练。"),("test_iterations","评估迭代；1/5000/10000/20000/30000会保存PSNR/SSIM。"),("save_iterations","保存可直接测试的Gaussian模型；训练结束也会保存。"),("checkpoint_iterations","保存可续训断点，含优化器状态。"),("start_checkpoint","从MODEL/ckpt/chkpnt10000.pth继续；需同一model_path。")],[3000,6360])

doc.add_heading("8. 测试：test.py",1)
para(doc,"功能：加载模型，渲染训练/测试投影，查询三维体并保存PNG、NPY、NIfTI和评估指标。")
table(doc,["参数","解释、作用与例子"],[("-m / model_path","训练输出目录。"),("-s / source_path","数据集目录；建议显式指定。"),("--iteration=-1","测试最新模型；也可写30000。"),("--skip_render_train","跳过训练投影渲染，节省时间。"),("--skip_render_test","跳过测试投影渲染。"),("--skip_recon","跳过三维体重建。"),("--quiet","减少输出。"),("--compute_cov3D_python / --debug","与训练保持一致；排错时使用。")],[3000,6360])

doc.add_heading("9. 旧/辅助脚本参数",1)
para(doc,"prepare_refcorr_r2.py：input_dir（原始TIFF）、sart_volume（SART体）、output_dir、n_train=50、pixel_subsample=4、volume_scale=1.0、sVoxel=75 75 75。用于已有SART体的RefCorr流程。")
para(doc,"data_generator/real_dataset/generate_data.py：data/output、proj_subsample=4、proj_rescale=400、object_scale=50、n_train=75、n_test=100、nVoxel=256 256 256、sVoxel=2 2 2、offOrigin、offDetector、accuracy=0.5。用于FIPS/MAT数据转换。")

doc.add_heading("10. SiO₂ 调参检查清单",1)
bullets(doc,["确认181张TIFF和config.txt角度一致；180°是重复端点。","原始探测器强度用intensity；已经是线积分才用line_integral。","中心必须在合理范围；宽度1028时约514附近，不能接受越界1523。","中心spread很大时先改用中间检测器行、检查shift_v，或使用scan_fbp_center.py。","prepare_fbp_tiff.py输出目录必须为空；scan_fbp_center.py输出父目录要先mkdir -p。","初始化点不足时降低density_thresh；显存不足时降低max_num_gaussians。","训练重点看eval/iter_005000和iter_030000的PSNR/SSIM及Gaussian数量。","test.py的iteration=-1表示最新模型。"])

doc.add_heading("11. 调参经验顺序",1)
table(doc,["现象","优先调整","示例"],[("中心越界/spread大","input_type、shift_v、中心算法/搜索范围","确认intensity；vo失败后相关性粗搜再精搜。"),("FBP噪声/振铃","filter、背景、投影缩放","hann；background=1；必要时关闭。"),("初始化点不够","density_thresh、n_points","0.05→0.02；50000→30000。"),("显存不足","max_num_gaussians、tv_vol_size","500000→300000；32→24。"),("过度平滑","lambda_tv、densify_until、grad阈值","0.05→0.02；15000→20000；5e-5→3e-5。"),("孤立噪点多","lambda_tv、grad阈值、density阈值","增大TV/梯度阈值，适度提高删除阈值。")],[2200,3300,3860])
para(doc,"说明：本手册按当前仓库暴露的命令行参数整理。库版本或硬件变化可能影响速度和数值，但不改变参数基本含义。")
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)

